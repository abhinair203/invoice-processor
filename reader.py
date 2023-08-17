from docx import Document
import os
import openpyxl

def generate_report():
     # Name of the output Excel file
    excel_file = "invoice_report.xlsx"
    # Creating a new Excel workbook and setting the active worksheet
    workbook = openpyxl.Workbook()
    worksheet = workbook.active
    # Adding headers to the Excel sheet
    worksheet.append(['Invoice ID', 'Total Quantity', 'Subtotal', 'Tax', 'Total'])
 # Loop through all files in the current directory
    for file_name in os.listdir('.'):
        if file_name.endswith('.docx'):  # Process only files that have a .docx extension
            doc = Document(file_name)

            # Extract Invoice ID
            invoice_id = file_name.replace(".docx", '')

            # Extract product quantities
            products = doc.paragraphs[1].text.split('\n')
            quantity = sum([int(product.split(':')[1]) for product in products if ':' in product])

            # Extract subtotal, tax, and total
            lines = doc.paragraphs[2].text.split('\n')
            subtotal = float(lines[1].split(':')[1])
            tax = float(lines[2].split(':')[1])
            total = subtotal + tax

            worksheet.append([invoice_id, quantity, subtotal, tax, total])

    workbook.save(excel_file)
    print(f"Report saved as {excel_file}")

generate_report()
